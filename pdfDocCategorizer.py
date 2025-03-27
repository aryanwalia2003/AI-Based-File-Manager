import os
import json
import shutil
import faiss
import numpy as np
import pdfplumber
import re
import base64
from PIL import Image
import google.generativeai as genai
import tempfile
import fitz  
from openai import OpenAI
from langchain.chat_models import AzureChatOpenAI
from langchain.schema import HumanMessage
import docx  

# API Keys and Configuration
AZURE_ENDPOINT = os.getenv("AZURE_ENDPOINT")
AZURE_API_KEY = os.getenv("AZURE_API_KEY")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# Azure OpenAI for Class Generation
llm = AzureChatOpenAI(
    azure_endpoint=AZURE_ENDPOINT,
    azure_deployment="gpt-4o",
    api_version="2023-09-01-preview",
    api_key=AZURE_API_KEY,
    temperature=0
)

# OpenAI for Embeddings
openai_client = OpenAI(api_key=OPENAI_API_KEY)

# Set up Gemini API
genai.configure(api_key=GOOGLE_API_KEY)
gemini_model = genai.GenerativeModel("gemini-1.5-flash")

def sanitize_folder_name(folder_name):
    """Remove invalid characters from folder names for Windows compatibility."""
    folder_name = folder_name.replace(" ", "_")  # Replace spaces with underscores
    folder_name = re.sub(r'[<>:"/\\|?*]', '', folder_name)  # Remove invalid characters
    folder_name = folder_name[:50]  # Limit folder name length to 50 characters
    return folder_name

def extract_text_from_pdf(pdf_path):
    """Extracts text from a PDF using pdfplumber (lightweight)."""
    text = ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
    except Exception as e:
        print(f"Error extracting text with pdfplumber from {pdf_path}: {e}")
    return text.strip()

def extract_text_from_docx(docx_path):
    """Extracts text from a DOCX file using python-docx."""
    try:
        # Open the Word document
        doc = docx.Document(docx_path)
        
        # Extract text from all paragraphs and tables
        full_text = []
        
        # Extract paragraph text
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Extract table text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        # Join paragraphs and table cells with newline
        return '\n'.join(full_text).strip()
    except Exception as e:
        print(f"Error extracting text from {docx_path}: {e}")
        return ""

def extract_pages_as_images(file_path, max_pages=3, is_pdf=True):
    """Extract the first few pages of a file as images."""
    image_paths = []
    temp_dir = tempfile.mkdtemp()
    
    try:
        if is_pdf:
            pdf_document = fitz.open(file_path)
            num_pages = min(len(pdf_document), max_pages)
            
            for page_num in range(num_pages):
                page = pdf_document[page_num]
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # Scale by 2 for better OCR
                image_path = os.path.join(temp_dir, f"page_{page_num}.png")
                pix.save(image_path)
                image_paths.append(image_path)
                
            pdf_document.close()
    except Exception as e:
        print(f"Error extracting pages as images: {e}")
    
    return image_paths, temp_dir

def extract_text_with_gemini(image_paths):
    """Extract text from images using Gemini's OCR capabilities."""
    print("Extracting text using Gemini OCR...")
    
    all_text = ""
    
    for image_path in image_paths:
        try:
            image = Image.open(image_path)
            response = gemini_model.generate_content([image])
            page_text = response.text
            all_text += page_text + "\n\n"
            
        except Exception as e:
            print(f"Error processing image {image_path} with Gemini: {e}")
    
    return all_text.strip()

def chunk_text(text, max_chunk_size=4000):
    """Splits text into chunks of approximately max_chunk_size characters."""
    if not text:
        return []
        
    chunks = []
    paragraphs = text.split('\n')
    current_chunk = ""
    
    for paragraph in paragraphs:
        if len(current_chunk) + len(paragraph) + 1 <= max_chunk_size:
            current_chunk += paragraph + "\n"
        else:
            if current_chunk:
                chunks.append(current_chunk.strip())
            current_chunk = paragraph + "\n"
    
    if current_chunk:
        chunks.append(current_chunk.strip())
    
    final_chunks = []
    for chunk in chunks:
        if len(chunk) > max_chunk_size:
            # Split by sentences if a paragraph is too long
            sentences = re.split(r'(?<=[.!?])\s+', chunk)
            temp_chunk = ""
            for sentence in sentences:
                if len(temp_chunk) + len(sentence) + 1 <= max_chunk_size:
                    temp_chunk += sentence + " "
                else:
                    if temp_chunk:
                        final_chunks.append(temp_chunk.strip())
                    temp_chunk = sentence + " "
            if temp_chunk:
                final_chunks.append(temp_chunk.strip())
        else:
            final_chunks.append(chunk)
    
    return final_chunks

def generate_embeddings(text):
    """Generates embeddings using OpenAI's text-embedding-ada-002 model.
    Handles large texts by chunking and averaging embeddings."""
    if not text:
        return np.zeros(1536)  # Return zero vector for empty text
    
    # Split text into manageable chunks
    chunks = chunk_text(text)
    
    if not chunks:
        return np.zeros(1536)
    
    # Get embeddings for each chunk
    chunk_embeddings = []
    for chunk in chunks:
        try:
            response = openai_client.embeddings.create(
                model="text-embedding-ada-002",
                input=chunk
            )
            chunk_embeddings.append(np.array(response.data[0].embedding))
        except Exception as e:
            print(f"Error generating embedding for chunk: {e}")
            # If error occurs, continue with other chunks
    
    # Return average embedding if we have any successful embeddings
    if chunk_embeddings:
        return np.mean(chunk_embeddings, axis=0)
    else:
        return np.zeros(1536)  # Fallback if all chunks failed

def generate_class_label(text):
    """Generates a class label for a document using Azure OpenAI (GPT-4o)."""
    # Use just the first chunk for classification
    chunks = chunk_text(text, max_chunk_size=2000)
    if not chunks:
        return "Unclassified"
        
    sample_text = chunks[0]
        
    prompt = f"Categorize the following document into a simple, one-word category (like Finance, Legal, Medical, Education, Resume, Marketing, etc.). Respond with ONLY the category name, nothing else:\n\n{sample_text}"
    try:
        response = llm.invoke([HumanMessage(content=prompt)])
        class_label = response.content.strip()
        # Ensure we get a simple category by taking just the first word if needed
        simple_label = class_label.split()[0] if len(class_label.split()) > 1 else class_label
        return simple_label
    except Exception as e:
        print(f"Error generating class label: {e}")
        return "Unclassified"

def process_documents(input_folder, output_folder):
    """Processes all PDFs and DOCX files: extracts text, generates embeddings, and clusters them."""
    os.makedirs(output_folder, exist_ok=True)
    
    # Find PDF and DOCX files
    document_files = [
        f for f in os.listdir(input_folder) 
        if f.lower().endswith((".pdf", ".docx"))
    ]
    
    if not document_files:
        print(f"No PDF or DOCX files found in {input_folder}")
        return
        
    index = faiss.IndexFlatL2(1536)  # FAISS index for 1536-dim embeddings
    doc_data = {}
    embeddings_list = []
    doc_paths = []  # Keep track of processed documents for index mapping

    print(f"Found {len(document_files)} document files to process")
    
    for i, doc_file in enumerate(document_files):
        print(f"Processing {i+1}/{len(document_files)}: {doc_file}")
        doc_path = os.path.join(input_folder, doc_file)
        
        try:
            # Determine file type and extract text
            if doc_file.lower().endswith(".pdf"):
                text = extract_text_from_pdf(doc_path)
                is_pdf = True
            elif doc_file.lower().endswith(".docx"):
                text = extract_text_from_docx(doc_path)
                is_pdf = False
            
            # If no text is extracted, use Gemini OCR for PDFs
            if not text and is_pdf:
                print(f"No text extracted from {doc_file} using pdfplumber. Trying Gemini OCR...")
                image_paths, temp_dir = extract_pages_as_images(doc_path, max_pages=3)
                
                if image_paths:
                    text = extract_text_with_gemini(image_paths)
                    
                    # Clean up temporary directory with images
                    shutil.rmtree(temp_dir, ignore_errors=True)
            
            # If still no text, skip this document
            if not text:
                print(f"Skipping {doc_file} (No text could be extracted)")
                continue
                
            embedding = generate_embeddings(text)
            embeddings_list.append(embedding)
            doc_paths.append(doc_file)
            
            class_label = generate_class_label(text)
            sanitized_class = sanitize_folder_name(class_label)  # Sanitize the class label for folder name

            doc_data[doc_file] = class_label
            
            # Create category folder and MOVE the file
            class_folder = os.path.join(output_folder, sanitized_class)
            os.makedirs(class_folder, exist_ok=True)
            
            destination_path = os.path.join(class_folder, doc_file)
            shutil.move(doc_path, destination_path)
            print(f"Moved {doc_file} -> Class: {class_label} (Folder: {sanitized_class})")
            
        except Exception as e:
            print(f"Error processing {doc_file}: {e}")
    
    # Add all embeddings to FAISS index
    if embeddings_list:
        embeddings_array = np.array(embeddings_list, dtype=np.float32)
        index.add(embeddings_array)
        
        # Save FAISS index and document metadata
        try:
            faiss.write_index(index, os.path.join(output_folder, "faiss_index.bin"))
            
            # Save document data with additional mapping information
            enhanced_doc_data = {
                "classifications": doc_data,
                "index_mapping": {i: doc_paths[i] for i in range(len(doc_paths))}
            }
            
            with open(os.path.join(output_folder, "document_metadata.json"), "w") as f:
                json.dump(enhanced_doc_data, f, indent=4)
                
            print(f"Successfully saved FAISS index and classification data")
        except Exception as e:
            print(f"Error saving index or classification data: {e}")
    else:
        print("No embeddings generated, skipping index creation")

if __name__ == "__main__":
    input_folder = r"C:\Users\Aryan Walia\OneDrive\Desktop\PDFTryFolder\docs"  # Update this
    output_folder = r"C:\Users\Aryan Walia\OneDrive\Desktop\PDFTryFolder\categorized"  # Update this
    process_documents(input_folder, output_folder)